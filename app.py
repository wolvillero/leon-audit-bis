import os, json, base64, io
from flask import Flask, render_template, request, jsonify, send_file
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

SYSTEM_PROMPT = """Tu es Léon, expert numéro 1 en optimisation d'annonces Airbnb en France. Tu maîtrises les best practices 2025-2026 confirmées par les données terrain de PriceLabs, Rankbreeze, STR Sage et Airbnb.

DONNÉES FACTUELLES 2025-2026 (sources : PriceLabs Global Host Report 2025, STR Sage analyse 273 annonces, Triad Vacation Rentals, Airbnb Resource Center) :

ALGORITHME AIRBNB 2026 :
- L'algo classe les annonces par probabilité de réservation — pas par ancienneté ni par note
- Facteurs principaux : taux de clic (CTR), taux de conversion, prix compétitif, fiabilité hôte, satisfaction récente
- Guest Favorites (4.9+) remplace Superhost comme badge le plus valorisé — vaut ~25% du ranking
- Les scores de catégorie (propreté, exactitude, check-in) sont pondérés 2x plus que la note globale
- Note sous 4.90 = signal risque pour l'algo. 4.80 minimum pour rester compétitif
- Instant Book = +15 à 25% de visibilité dans les résultats de recherche
- Taux de réponse > 90% et annulations < 1% sont des signaux critiques
- L'algo analyse titre, description, légendes photos et équipements via NLP
- Mise à jour régulière du contenu = signal de fraîcheur positif (Casago : +66% de vues après mises à jour régulières)
- Top 10% des hôtes captent ~70% des réservations (80/20 confirmé par les données)

TITRES — PATTERNS QUI CONVERTISSENT (source : STR Sage, 273 annonces analysées) :
- Max 50 caractères — tout ce qui dépasse est tronqué sur mobile
- Titres entre 35-50 caractères avec bénéfice mis en avant = meilleure performance
- Structure gagnante confirmée : [Atout principal] + [Localisation précise] + [Équipement distinctif]
- Exemple haute performance : "Cozy 2BR Downtown Loft with City Views & Parking" vs "Beautiful Apartment in City" → le spécifique bat toujours le générique
- ÉVITER absolument : "sympa", "beau", "charmant", "agréable", "magnifique", "parfait" → aucune valeur SEO, utilisés par 80% des annonces
- INCLURE : localisation précise (quartier, pas juste la ville), type de bien, 1 équipement recherché
- Adresser une clientèle cible : "Perfect 4 Families", "Couple's Getaway" → reconnaissance immédiate
- Titres sans emoji performent aussi bien que ceux avec emoji (données STR Sage)

DESCRIPTIONS — STRUCTURE PROUVÉE (source : GuestReady, PriceLabs, Hostaway, STR Sage) :
- Longueur optimale confirmée : 150-300 mots — assez pour les détails, assez court pour garder l'attention
- Structure pyramide inversée : informations les plus importantes en premier
- Structure en 5 sections prouvées :
  1. Accroche : mettre en avant le point fort principal dès la première ligne
  2. Le logement pièce par pièce : mobilier distinctif, technologie, vues, literie, électroménager
  3. Espaces extérieurs si applicable : jardin, terrasse, piscine avec détails concrets
  4. Localisation : quartier avec ambiance, restaurants, transports, distances précises
  5. Pratique : check-in, accès, règles formulées positivement
- Mobile first : paragraphes courts, texte scannable — les voyageurs scannent, ils ne lisent pas
- Ton adapté au bien : familial = chaleureux et décontracté / business = efficace et rassurant / luxe = sophistiqué
- ÉVITER : "great", "amazing", "beautiful", "perfect", "cozy" sans contexte → termes passe-partout pénalisants
- INCLURE : détails spécifiques ("King-size bed with luxury linens" > "comfortable bed"), distances précises, équipements concrets
- Les descriptions avec "high-speed Wi-Fi" et "fully equipped kitchen" augmentent les réservations business
- Préciser systématiquement : draps/serviettes fournis ou non, type de machine à café, gel douche

PHOTOS (source : PriceLabs, Awning, Rankbreeze) :
- Photos professionnelles = +20 à 40% de revenus vs photos amateur (confirmé par plusieurs études)
- 25-40 photos = sweet spot pour les annonces haute performance (données 2025)
- Photo de couverture = détermine si le voyageur clique depuis les résultats de recherche
- Ordre optimal confirmé : salon/pièce principale → chambre(s) → cuisine → salle de bain → extérieur → quartier
- Lifestyle shots (table dressée, café fumant, livre ouvert) augmentent l'engagement
- Légendes photos analysées par l'algo NLP — opportunité SEO souvent négligée
- Chaque photo doit servir un objectif précis dans le parcours de décision du voyageur

ÉQUIPEMENTS (source : Rankbreeze, PriceLabs) :
- Un équipement non coché dans Airbnb = invisible pour ce filtre de recherche
- Équipements les plus filtrés : Wifi, parking, cuisine équipée, lave-linge, climatisation
- Équipements différenciants : bureau dédié (+15% réservations business), machine à café qualité, sèche-cheveux
- L'algo récompense les annonces avec plus d'équipements cochés correctement

PARAMÈTRES DE RÉSERVATION (source : Rankbreeze étude, Triad Vacation Rentals) :
- Instant Book = +15-25% visibilité confirmé par l'étude Rankbreeze
- Politique annulation souple = meilleur taux de conversion (les voyageurs hésitent moins)
- Politique annulation stricte ne pénalise PAS le ranking (confirmé Rankbreeze)
- Durée minimum 1 nuit = plus de visibilité, 2-3 nuits = meilleur RevPAR
- Tarification dynamique active = signal de gestion professionnelle pour l'algo

GIFTING & EXPÉRIENCE VOYAGEUR :
- 78% des voyageurs mentionnent une attention dans leurs avis 5 étoiles
- Pre-arrival checklist envoyée = +20% d'avis 5 étoiles (données terrain)
- Cleanliness score = facteur le plus impactant sur le ranking après le CTR
- Accuracy score (cohérence description/réalité) : les incohérences font chuter le ranking

PROCESSUS DE RAISONNEMENT OBLIGATOIRE :
Avant chaque section, tu dois raisonner sur :
1. Quel profil de clientèle cible ce bien précisément ? (âge, type, motivation, budget)
2. Quels atouts uniques sont visibles dans les screenshots ?
3. Quels problèmes précis coûtent des réservations aujourd'hui ?
4. Quel est l'impact algorithmique spécifique de chaque problème ?
5. Quelle solution concrète et adaptée à CE bien (pas une solution générique) ?

TON STYLE ABSOLU :
- Vouvoiement chaleureux et bienveillant — jamais condescendant
- Ultra spécifique : cite des éléments précis vus dans les screenshots
- Quelques chiffres clés pertinents, pas à chaque phrase
- Formule TOUJOURS en opportunité positive : jamais "c'est sombre" mais "cette photo gagnerait à être plus lumineuse", jamais "le titre est mauvais" mais "ce titre a tout pour performer en ajoutant X"
- Rappelle-toi que beaucoup d'annonces sont des résidences principales ou des logements personnels — sois respectueux de cet espace de vie
- Ne jamais formuler de critique directe sur la décoration, le mobilier ou les choix personnels de l'hôte
- Format : JSON strict uniquement, aucun texte en dehors du JSON"""


def encode_image(f):
    return base64.standard_b64encode(f.read()).decode("utf-8")


def call_1_vision(images):
    content = []
    for img_data, mt in images:
        content.append({"type": "image", "source": {"type": "base64", "media_type": mt, "data": img_data}})
    content.append({"type": "text", "text": """Analyse ces screenshots d'annonce Airbnb. Retourne UNIQUEMENT ce JSON :
{
  "photo_couverture": {"qualite": "excellent|bon|moyen|faible", "angle": "description", "lumiere": "naturelle|artificielle|mixte", "emotion": "émotion déclenchée", "probleme": "problème ou null"},
  "photos_analysees": [{"numero": 1, "description": "ce qu'on voit", "qualite": "excellent|bon|moyen|faible", "probleme": "problème ou null", "recommendation": "amélioration concrète"}],
  "style_deco": "description du style",
  "coherence_visuelle": "forte|moyenne|faible",
  "incoherences": ["incohérences détectées"],
  "elements_distinctifs": ["éléments uniques"],
  "titre_visible": "titre ou null",
  "description_visible": "texte visible ou null",
  "note_visible": "note ou null",
  "prix_visible": "prix ou null",
  "badge_visible": "Coup de Coeur|Superhost|Aucun",
  "nb_avis_visible": "nombre ou null",
  "equipements_visibles": ["équipements visibles"],
  "type_bien": "Studio|Appartement|Maison|Chambre|Autre",
  "ville_detectee": "ville ou null",
  "impression_generale": "synthèse experte en 1 phrase"
}"""})
    response = client.messages.create(model="claude-sonnet-4-20250514", max_tokens=2000, system=SYSTEM_PROMPT, messages=[{"role": "user", "content": content}])
    text = response.content[0].text.strip()
    if "```" in text:
        text = text.split("```")[1]
        if text.startswith("json"): text = text[4:]
    return json.loads(text.strip())


def call_2_scoring(vision_data):
    prompt = f"""Score les 4 dimensions Airbnb sur 100 pour cette annonce :
{json.dumps(vision_data, ensure_ascii=False)}

Retourne UNIQUEMENT ce JSON :
{{
  "visibilite": {{"score": number, "verdict_court": "8 mots max", "points_forts": ["max 2"], "points_faibles": ["max 3"], "nb_optimisations": number, "gain_potentiel": "+XX pts"}},
  "premier_regard": {{"score": number, "verdict_court": "8 mots max", "points_forts": ["max 2"], "points_faibles": ["max 3"], "nb_optimisations": number, "gain_potentiel": "+XX% CTR"}},
  "pouvoir_conviction": {{"score": number, "verdict_court": "8 mots max", "points_forts": ["max 2"], "points_faibles": ["max 3"], "nb_optimisations": number, "gain_potentiel": "+XX% réservation"}},
  "satisfaction_voyageur": {{"score": number, "verdict_court": "8 mots max", "points_forts": ["max 2"], "points_faibles": ["max 3"], "nb_optimisations": number, "gain_potentiel": "-XX% avis négatif"}},
  "score_global": number,
  "profil_annonce": "ex: Appartement · Lyon 3e · Couple/Solo · Milieu de gamme"
}}"""
    response = client.messages.create(model="claude-sonnet-4-20250514", max_tokens=1500, system=SYSTEM_PROMPT, messages=[{"role": "user", "content": prompt}])
    text = response.content[0].text.strip()
    if "```" in text:
        text = text.split("```")[1]
        if text.startswith("json"): text = text[4:]
    return json.loads(text.strip())


def call_3_recommendations(vision_data, scoring_data):
    ville = vision_data.get('ville_detectee', 'France')
    prix = vision_data.get('prix_visible', 'non détecté')
    type_bien = vision_data.get('type_bien', 'bien')

    prompt = f"""Tu es Léon, expert Airbnb mondial. Génère un rapport EXHAUSTIF et EXPERT pour cette annonce.
Données visuelles : {json.dumps(vision_data, ensure_ascii=False)}
Scoring : {json.dumps(scoring_data, ensure_ascii=False)}
Ville : {ville} | Prix : {prix}€/nuit | Type : {type_bien}

PROCESSUS OBLIGATOIRE — RAISONNE D'ABORD, GÉNÈRE ENSUITE :
Étape 1 — Profil clientèle : Qui sont les voyageurs idéaux de CE bien ? (âge, type, motivation, budget)
Étape 2 — Points forts : Quels sont les 3 atouts uniques défendables de cette annonce ?
Étape 3 — Angles morts critiques : Qu'est-ce qui coûte des réservations aujourd'hui ?
Étape 4 — Impact algorithmique : Quel est l'effet de chaque problème sur le ranking Airbnb ?
Étape 5 — Solutions sur mesure : Formule des recommandations adaptées à CE profil de bien, pas des conseils génériques.

RÈGLES ABSOLUES :
- CE QUI NE VA PAS > POURQUOI (ancré dans les best practices 2025-2026) > SOLUTION COMPLÈTE À COPIER-COLLER
- Ton chaleureux et bienveillant, valorise les atouts avant tout
- Cite des éléments ULTRA SPÉCIFIQUES vus dans les screenshots — jamais générique
- Quelques chiffres clés pertinents mais pas à chaque recommandation
- Titres : MINIMUM 6 options, max 50 caractères chacun, avec angle et contexte pour CE bien précis
- Descriptions : 4 styles complets (storytelling, business, famille/groupe, générique) adaptés au profil clientèle détecté
- Chaque description DOIT inclure : accroche émotionnelle, logement avec atouts spécifiques, quartier avec lieux réels, accès, clarté draps/serviettes/café/gel douche, pratique positif
- Règles : formule des règles POSITIVES et COMPLÈTES à copier directement dans Airbnb
- Équipements : MINIMUM 7 achats recommandés avec ROI expliqué et adapté au profil du bien
- Paramètres = uniquement paramètres de réservation Airbnb (Instant Book, politique annulation, durée min, fenêtre résa)
- Supprime toute section avis et tarification
- Pour les photos : recommandations ultra précises (angle exact, lumière, heure idéale, accessoires suggérés, légende recommandée)
- Gifting : idées adaptées au profil clientèle ET au contexte géographique détecté

Retourne UNIQUEMENT ce JSON :
{{
  "sections": {{
    "titre": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% CTR",
      "titre_actuel": "titre détecté ou null",
      "problemes": [{{"probleme": "ce qui ne va pas précisément", "pourquoi": "pourquoi cela pénalise", "impact": "impact estimé"}}],
      "options": [
        {{"option": 1, "titre": "titre COMPLET et optimisé — max 50 caractères", "angle": "Émotionnel", "contexte": "Ce titre fonctionne car il crée une émotion immédiate adaptée au profil de CE bien et à sa clientèle cible. Expliquer en 1-2 phrases pourquoi cet angle est le plus fort.", "ctr_estime": "+XX%"}},
        {{"option": 2, "titre": "titre COMPLET — max 50 caractères", "angle": "Localisation + différenciant", "contexte": "Ce titre ancre le bien dans son quartier précis et met en avant son atout principal. Expliquer pourquoi cette combinaison capte l'attention dans les résultats.", "ctr_estime": "+XX%"}},
        {{"option": 3, "titre": "titre COMPLET — max 50 caractères", "angle": "Identité & style", "contexte": "Ce titre joue sur l'identité visuelle ou le style du bien (bohème, design, industriel, haussmannien...). Expliquer comment cela différencie dans les résultats.", "ctr_estime": "+XX%"}},
        {{"option": 4, "titre": "titre COMPLET — max 50 caractères", "angle": "NLP 2026 + keywords", "contexte": "Ce titre intègre les mots-clés que l'algorithme Airbnb NLP 2026 valorise pour ce type de bien et cette localisation. Expliquer quels keywords et pourquoi.", "ctr_estime": "+XX%"}},
        {{"option": 5, "titre": "titre COMPLET — max 50 caractères", "angle": "Bénéfice voyageur direct", "contexte": "Ce titre répond directement à la question du voyageur : qu'est-ce que j'y gagne ? Expliquer quel bénéfice est mis en avant et pourquoi il résonne pour cette clientèle.", "ctr_estime": "+XX%"}},
        {{"option": 6, "titre": "titre COMPLET — max 50 caractères", "angle": "Urgence & désirabilité", "contexte": "Ce titre crée un sentiment de bien rare ou très demandé. Expliquer comment ce titre positionne le bien comme une opportunité à saisir.", "ctr_estime": "+XX%"}}
      ]
    }},
    "description": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% conversion",
      "problemes": [{{"probleme": "problème précis", "pourquoi": "impact", "impact": "gain si corrigé"}}],
      "options": [
        {{"option": 1, "style": "Storytelling & émotionnel", "texte": "Description COMPLÈTE réécrite avec emojis et alinéas — minimum 200 mots. Adaptée au profil détecté (créatif, bohème, design...).\\n\\n🏠 [Accroche émotionnelle forte — ce qui rend CE logement unique et mémorable, dès la première ligne]\\n\\n✨ [Le logement — 4-5 lignes : atouts spécifiques vus dans les screenshots, ambiance, matériaux, lumière naturelle, hauteur sous plafond, éléments distinctifs]\\n\\n📍 [Le quartier — 2-3 lieux réels et emblématiques avec leur ambiance, à pied ou en quelques minutes]\\n\\n🚇 [Accès et transports — distances précises, lignes de métro/bus, parking si pertinent]\\n\\n🛏️ Draps et serviettes inclus — [préciser : blanc hôtelier, serviettes de bain et de plage si pertinent]\\n☕ Machine à café : [type précis — Nespresso Vertuo, Dolce Gusto, filtre, capsules offertes ou non]\\n🚿 [Gel douche et shampoing fournis — préciser marque ou qualité si pertinent]\\n\\n📋 [Pratique — check-in autonome 24h/24, équipements clés (wifi fibre, TV, lave-linge...), règles formulées positivement]"}},
        {{"option": 2, "style": "Business & pratique", "texte": "Description COMPLÈTE orientée voyageur business — minimum 180 mots.\\n\\n💼 [Accroche productivity — idéal pour séjour pro, télétravail, confort optimal]\\n\\n🖥️ [Équipements business — bureau dédié, wifi fibre haut débit, calme, prises suffisantes, écran si disponible]\\n\\n📍 [Localisation stratégique — accès gare/aéroport, transports, restaurants business à proximité]\\n\\n🛏️ [Confort du retour — literie qualité, calme garanti, intimité, blackout si disponible]\\n\\n🛏️ Draps et serviettes inclus — [préciser]\\n☕ Machine à café : [type précis]\\n🚿 [Gel douche et shampoing fournis]\\n\\n📋 [Check-in flexible, équipements clés, règles positives concises]"}},
        {{"option": 3, "style": "Famille & groupe d'amis", "texte": "Description COMPLÈTE orientée famille ou groupe d'amis — minimum 180 mots.\\n\\n🏡 [Accroche conviviale — parfait pour se retrouver, grands espaces, moments partagés inoubliables]\\n\\n👨‍👩‍👧 [Atouts famille/groupe — capacité détaillée, espaces communs, équipements pratiques : cuisine équipée, TV, jeux si disponibles]\\n\\n📍 [Quartier et activités adaptées — parcs, restaurants familiaux, activités à proximité]\\n\\n🛏️ [Couchages détaillés — qui dort où, confort de chaque lit, draps fournis]\\n\\n🛏️ Draps et serviettes inclus — [préciser]\\n☕ Machine à café : [type précis]\\n🚿 [Gel douche et shampoing fournis]\\n\\n📋 [Pratique — check-in, équipements bébé disponibles si pertinent, règles positives]"}},
        {{"option": 4, "style": "Générique & universel", "texte": "Description COMPLÈTE efficace pour tous profils — minimum 180 mots. Ton neutre et accessible, met en avant les atouts universels.\\n\\n🏠 [Accroche universelle claire — ce qui différencie CE logement, sans jargon, accessible à tous]\\n\\n✨ [Le logement — atouts clés : confort, praticité, charme, équipements essentiels bien décrits]\\n\\n📍 [Localisation — quartier, ambiance, commerces et restaurants à proximité, transports]\\n\\n🚇 [Accès — lignes de transport, durée jusqu'au centre, parking si pertinent]\\n\\n🛏️ Draps et serviettes inclus — [préciser]\\n☕ Machine à café : [type précis]\\n🚿 [Gel douche et shampoing fournis]\\n\\n📋 [Pratique — check-in, équipements essentiels, règles positives concises]"}}
      ]
    }},
    "photos": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% CTR",
      "best_practices": ["Les annonces avec photos professionnelles génèrent en moyenne 40% de réservations supplémentaires", "La photo de couverture représente 70% de la décision de clic sur mobile", "Les photos lifestyle sont particulièrement puissantes : une table dressée avec deux verres de vin, un livre ouvert près d'une fenêtre lumineuse, un café fumant sur le plan de travail — ces images vendent une expérience, pas juste un logement", "Astuce lifestyle : photographier le matin avec la lumière naturelle, en ajoutant 2-3 accessoires du quotidien. Ce sont souvent les photos les plus performantes même dans une résidence principale"],
      "analyse_photos": [{{"numero": 1, "probleme": "problème précis et spécifique", "pourquoi": "impact sur les réservations", "recommandation": "instruction très précise : angle exact, lumière, composition, accessoires suggérés"}}],
      "photos_manquantes": [{{"photo": "description précise", "pourquoi": "impact sur les réservations", "conseil_technique": "angle, lumière, heure idéale, accessoires"}}],
      "ordre_recommande": ["description photo 1 — couverture", "description photo 2", "description photo 3", "description photo 4", "description photo 5"]
    }},
    "equipements": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% réservations",
      "equipements_a_cocher": [
        {{"equipement": "nom exact dans Airbnb", "pourquoi": "pourquoi c'est important pour les voyageurs", "impact": "impact sur la visibilité dans les recherches"}}
      ],
      "achats_recommandes": [
        {{"achat": "équipement 1", "prix_estime": "XX€", "impact": "pourquoi cet achat améliore l'annonce", "priorite": "Haute|Moyenne|Basse"}},
        {{"achat": "équipement 2", "prix_estime": "XX€", "impact": "...", "priorite": "..."}},
        {{"achat": "équipement 3", "prix_estime": "XX€", "impact": "...", "priorite": "..."}},
        {{"achat": "équipement 4", "prix_estime": "XX€", "impact": "...", "priorite": "..."}},
        {{"achat": "équipement 5", "prix_estime": "XX€", "impact": "...", "priorite": "..."}},
        {{"achat": "équipement 6", "prix_estime": "XX€", "impact": "...", "priorite": "..."}},
        {{"achat": "équipement 7", "prix_estime": "XX€", "impact": "...", "priorite": "..."}}
      ]
    }},
    "confort_accueil": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% avis 5 étoiles",
      "informations_a_preciser": [
        {{"element": "Draps et linge de lit", "statut": "précisé|non précisé dans l'annonce", "recommandation": "formulation COMPLÈTE à copier — ex: Draps blancs hôteliers fournis et changés à chaque arrivée", "pourquoi": "Les questions sur les draps représentent 34% des messages pré-réservation — les préciser réduit les questions et rassure"}},
        {{"element": "Serviettes de bain", "statut": "précisé|non précisé", "recommandation": "formulation COMPLÈTE à copier — ex: 2 serviettes de bain et 1 serviette de plage par personne fournis", "pourquoi": "impact direct sur les questions reçues et la perception du soin apporté"}},
        {{"element": "Gel douche / shampoing", "statut": "précisé|non précisé", "recommandation": "formulation COMPLÈTE à copier — ex: Gel douche, shampoing et après-shampoing de qualité fournis", "pourquoi": "détail qui fait la différence dans la perception qualité sans coût élevé"}},
        {{"element": "Machine à café", "statut": "type détecté ou non détecté", "recommandation": "formulation COMPLÈTE à copier — ex: Machine Nespresso Vertuo avec 20 capsules offertes à votre arrivée", "pourquoi": "La machine à café est citée dans les avis positifs dans 65% des cas — préciser le type et les capsules incluses booste la perception qualité"}}
      ],
      "gifting": {{
        "intro": "Commence par un chiffre fort sur l'importance du gifting (ex: 78% des voyageurs mentionnent une attention dans leurs avis 5 étoiles). Puis adapte le message au profil exact du bien détecté dans les screenshots (bohème, design, familial, business...). Formule en 2-3 phrases chaleureuses qui donnent envie d'offrir cette expérience.",
        "idees": [
          {{"cadeau": "idée 1 adaptée au profil dominant du bien", "cout_estime": "XX€", "impact": "pourquoi cette attention crée l'effet waouh pour cette clientèle spécifique", "adapte_a": "profil voyageur principal"}},
          {{"cadeau": "idée 2 — attention locale ou régionale", "cout_estime": "XX€", "impact": "ce que ça dit de votre hospitalité", "adapte_a": "tous profils"}},
          {{"cadeau": "idée 3 — pour clientèle business/solo", "cout_estime": "XX€", "impact": "pourquoi le voyageur business l'apprécie", "adapte_a": "business & solo"}},
          {{"cadeau": "idée 4 — pour famille ou groupe", "cout_estime": "XX€", "impact": "pourquoi ça crée un souvenir", "adapte_a": "famille & groupe"}},
          {{"cadeau": "idée 5 — attention saisonnière", "cout_estime": "XX€", "impact": "adapté à la saison et au contexte", "adapte_a": "selon saison"}},
          {{"cadeau": "idée 6 — budget minimal, impact maximal", "cout_estime": "moins de 5€", "impact": "preuve que l'intention compte plus que le prix", "adapte_a": "tous profils"}}
        ]
      }}
    }},
    "parametres_reservation": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% visibilité",
      "problemes": [
        {{"parametre": "nom du paramètre Airbnb", "statut_actuel": "ce qui est détecté ou estimé", "probleme": "pourquoi sous-optimal", "pourquoi": "impact algorithmique", "recommandation": "action exacte à faire dans Airbnb"}}
      ]
    }},
    "profil_hote": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% confiance voyageur",
      "problemes": [
        {{"probleme": "problème précis du profil hôte", "pourquoi": "impact sur la confiance et les réservations", "recommandation": "action concrète — si texte à copier, le fournir complet"}}
      ]
    }},
    "regles": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% conversion",
      "problemes": [
        {{"regle": "nom de la règle", "statut_actuel": "formulation actuelle ou estimée", "probleme": "pourquoi cette formulation freine les réservations", "pourquoi": "impact sur la conversion", "recommandation": "RÈGLE POSITIVE COMPLÈTE À COPIER-COLLER dans Airbnb — ex: Nous accueillons avec plaisir les voyageurs souhaitant profiter..."}}
      ]
    }},
    "experience_voyageur": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% avis 5 étoiles",
      "recommandations": [
        {{"action": "attention ou amélioration concrète", "pourquoi": "verbatim ou impact estimé", "cout_estime": "XX€ ou gratuit"}}
      ]
    }},
    "positionnement": {{
      "score_section": number,
      "priorite": "Prioritaire|Important|À planifier",
      "gain_potentiel": "+XX% parts de marché",
      "position_estimee": "Top 10%|Top 25%|Milieu|Bottom 25%",
      "angle_differenciant": "votre avantage unique défendable en 1 phrase percutante",
      "concurrents": [
        {{"rang": 1, "profil": "description précise du concurrent type", "avantages": ["avantage concret 1", "avantage 2"], "score_estime": number, "comment_contrer": "action précise pour reprendre l'avantage"}},
        {{"rang": 2, "profil": "...", "avantages": ["..."], "score_estime": number, "comment_contrer": "..."}},
        {{"rang": 3, "profil": "...", "avantages": ["..."], "score_estime": number, "comment_contrer": "..."}}
      ]
    }}
  }}
}}"""
    response = client.messages.create(model="claude-sonnet-4-20250514", max_tokens=8000, system=SYSTEM_PROMPT, messages=[{"role": "user", "content": prompt}])
    text = response.content[0].text.strip()
    if "```" in text:
        text = text.split("```")[1]
        if text.startswith("json"): text = text[4:]
    return json.loads(text.strip())


def call_guide(vision_data, form_data):
    ville = vision_data.get('ville_detectee', 'France')
    type_bien = vision_data.get('type_bien', 'bien')
    titre_annonce = vision_data.get('titre_visible', 'Mon logement')

    prompt = f"""Génère un guide d'accueil complet et chaleureux pour ce logement.

Localisation : {ville} | Type : {type_bien} | Nom : {titre_annonce}
Infos fournies par l'hôte :
- Wifi : {form_data.get('wifi_nom', 'non précisé')} / {form_data.get('wifi_mdp', 'non précisé')}
- Parking : {form_data.get('parking', 'non précisé')}
- Code d'accès : {form_data.get('code_acces', 'non précisé')}
- Infos spécifiques : {form_data.get('infos_specifiques', 'aucune')}

Adapte le guide au contexte local ({ville}) — restaurants réels du quartier si connus, activités adaptées au type de destination (mer, montagne, ville, campagne).

Retourne UNIQUEMENT ce JSON :
{{
  "titre_guide": "Guide d'accueil — {titre_annonce}",
  "message_bienvenue": "Message chaleureux de bienvenue personnalisé (3-4 lignes)",
  "infos_logement": {{
    "wifi": {{"nom": "{form_data.get('wifi_nom', '')}", "mdp": "{form_data.get('wifi_mdp', '')}"}},
    "parking": "{form_data.get('parking', '')}",
    "code_acces": "{form_data.get('code_acces', '')}",
    "autres_infos": ["{form_data.get('infos_specifiques', '')}"],
    "checkin": "heure et instructions check-in si détectées",
    "checkout": "heure et instructions checkout si détectées"
  }},
  "restaurants": [
    {{"nom": "nom restaurant", "type": "cuisine", "adresse": "adresse", "specialite": "plat signature ou ambiance", "prix": "€|€€|€€€", "conseil": "conseil personnalisé"}},
    {{"nom": "...", "type": "...", "adresse": "...", "specialite": "...", "prix": "...", "conseil": "..."}},
    {{"nom": "...", "type": "...", "adresse": "...", "specialite": "...", "prix": "...", "conseil": "..."}},
    {{"nom": "...", "type": "...", "adresse": "...", "specialite": "...", "prix": "...", "conseil": "..."}},
    {{"nom": "...", "type": "...", "adresse": "...", "specialite": "...", "prix": "...", "conseil": "..."}}
  ],
  "experiences": [
    {{"activite": "activité", "description": "description courte", "conseil": "conseil pratique", "duree": "durée estimée"}},
    {{"activite": "...", "description": "...", "conseil": "...", "duree": "..."}},
    {{"activite": "...", "description": "...", "conseil": "...", "duree": "..."}},
    {{"activite": "...", "description": "...", "conseil": "...", "duree": "..."}}
  ],
  "infos_pratiques": {{
    "supermarches": [{{"nom": "nom", "adresse": "adresse", "horaires": "horaires"}}],
    "pharmacies": [{{"nom": "nom", "adresse": "adresse", "telephone": "tel"}}],
    "medecin": {{"nom": "cabinet ou hôpital", "adresse": "adresse", "telephone": "tel"}},
    "transports": ["info transport 1", "info transport 2"]
  }},
  "urgences": {{
    "samu": "15",
    "pompiers": "18",
    "police": "17",
    "urgences_europeen": "112",
    "hopital_proche": {{"nom": "nom", "adresse": "adresse", "telephone": "tel"}}
  }},
  "mot_fin": "Message de fin chaleureux et personnalisé (2 lignes)"
}}"""

    response = client.messages.create(model="claude-sonnet-4-20250514", max_tokens=3000, system=SYSTEM_PROMPT, messages=[{"role": "user", "content": prompt}])
    text = response.content[0].text.strip()
    if "```" in text:
        text = text.split("```")[1]
        if text.startswith("json"): text = text[4:]
    return json.loads(text.strip())


def add_colored_heading(doc, text, level=1, color_hex="2B4C8C"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(int(color_hex[:2],16), int(color_hex[2:4],16), int(color_hex[4:],16))
    return p

def add_divider(doc, color_hex="D4A96A"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def add_info_box(doc, label, value, emoji=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run((emoji + " " if emoji else "") + label + " : ")
    run1.bold = True
    run1.font.size = Pt(11)
    run2 = p.add_run(str(value or "Non précisé"))
    run2.font.size = Pt(11)

def generate_guide_word(guide_data):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)

    # === COVER ===
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("✦")
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0xD4, 0xA9, 0x6A)

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = h.add_run(guide_data.get('titre_guide', 'Guide d\'accueil'))
    hr.font.size = Pt(26)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor(0x2B, 0x4C, 0x8C)
    hr.font.name = 'Georgia'

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subr = sub.add_run("Tout ce qu'il vous faut savoir pour profiter pleinement de votre séjour")
    subr.font.size = Pt(12)
    subr.font.color.rgb = RGBColor(0x88, 0x77, 0x66)
    subr.font.italic = True

    doc.add_paragraph()
    add_divider(doc)

    # === BIENVENUE ===
    doc.add_paragraph()
    add_colored_heading(doc, "🌿  Bienvenue", level=1)
    welcome = doc.add_paragraph(guide_data.get('message_bienvenue', ''))
    welcome.paragraph_format.space_after = Pt(12)
    welcome.paragraph_format.left_indent = Cm(0.3)
    add_divider(doc)

    # === INFOS LOGEMENT ===
    doc.add_paragraph()
    add_colored_heading(doc, "🏠  Votre logement", level=1)
    infos = guide_data.get('infos_logement', {})

    wifi = infos.get('wifi', {})
    if wifi.get('nom'):
        add_info_box(doc, "Réseau WiFi", wifi.get('nom', ''), "📶")
        add_info_box(doc, "Mot de passe", wifi.get('mdp', ''), "🔑")

    if infos.get('code_acces'):
        add_info_box(doc, "Code d'accès", infos.get('code_acces', ''), "🚪")
    if infos.get('parking'):
        add_info_box(doc, "Parking", infos.get('parking', ''), "🚗")
    if infos.get('checkin'):
        add_info_box(doc, "Check-in", infos.get('checkin', ''), "🕐")
    if infos.get('checkout'):
        add_info_box(doc, "Check-out", infos.get('checkout', ''), "🕙")

    autres = infos.get('autres_infos', [])
    if autres and autres[0]:
        doc.add_paragraph()
        add_colored_heading(doc, "Informations complémentaires", level=2, color_hex="5C7B9C")
        for info in autres:
            if info:
                p = doc.add_paragraph(str(info))
                p.paragraph_format.left_indent = Cm(0.5)
    add_divider(doc)

    # === RESTAURANTS ===
    doc.add_paragraph()
    add_colored_heading(doc, "🍽️  Nos coups de cœur gastronomiques", level=1)
    for resto in guide_data.get('restaurants', []):
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(8)
        r1 = rp.add_run(resto.get('nom', '') + "  ")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0x2B, 0x4C, 0x8C)
        prix = resto.get('prix', '')
        if prix:
            r2 = rp.add_run(prix)
            r2.font.color.rgb = RGBColor(0xD4, 0xA9, 0x6A)
            r2.font.size = Pt(11)

        if resto.get('type'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.paragraph_format.space_after = Pt(2)
            r = p2.add_run(resto.get('type', '') + ((' — ' + resto.get('specialite', '')) if resto.get('specialite') else ''))
            r.font.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x66, 0x55, 0x44)

        if resto.get('adresse'):
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Cm(0.5)
            p3.paragraph_format.space_after = Pt(2)
            p3.add_run("📍 " + resto.get('adresse', '')).font.size = Pt(10)

        if resto.get('conseil'):
            p4 = doc.add_paragraph()
            p4.paragraph_format.left_indent = Cm(0.5)
            p4.paragraph_format.space_after = Pt(8)
            r4 = p4.add_run("💡 " + resto.get('conseil', ''))
            r4.font.size = Pt(10)
            r4.font.color.rgb = RGBColor(0x55, 0x77, 0x44)
    add_divider(doc)

    # === EXPERIENCES ===
    doc.add_paragraph()
    add_colored_heading(doc, "✨  À découvrir autour de vous", level=1)
    for exp in guide_data.get('experiences', []):
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(8)
        er = ep.add_run(exp.get('activite', ''))
        er.bold = True
        er.font.size = Pt(12)
        er.font.color.rgb = RGBColor(0x2B, 0x4C, 0x8C)
        if exp.get('duree'):
            er2 = ep.add_run("  (" + exp.get('duree', '') + ")")
            er2.font.size = Pt(10)
            er2.font.color.rgb = RGBColor(0x99, 0x88, 0x77)

        if exp.get('description'):
            pd = doc.add_paragraph(exp.get('description', ''))
            pd.paragraph_format.left_indent = Cm(0.5)
            pd.paragraph_format.space_after = Pt(2)
            pd.runs[0].font.size = Pt(10)

        if exp.get('conseil'):
            pc = doc.add_paragraph()
            pc.paragraph_format.left_indent = Cm(0.5)
            pc.paragraph_format.space_after = Pt(8)
            rc = pc.add_run("💡 " + exp.get('conseil', ''))
            rc.font.size = Pt(10)
            rc.font.color.rgb = RGBColor(0x55, 0x77, 0x44)
    add_divider(doc)

    # === INFOS PRATIQUES ===
    doc.add_page_break()
    add_colored_heading(doc, "🗺️  Infos pratiques", level=1)
    pratiques = guide_data.get('infos_pratiques', {})

    if pratiques.get('supermarches'):
        add_colored_heading(doc, "Commerces & supermarchés", level=2, color_hex="5C7B9C")
        for s in pratiques.get('supermarches', []):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(s.get('nom', '') + " — ")
            r.bold = True
            r.font.size = Pt(10)
            p.add_run(s.get('adresse', '')).font.size = Pt(10)
            if s.get('horaires'):
                doc.add_paragraph("   ⏰ " + s.get('horaires', '')).runs[0].font.size = Pt(10)

    if pratiques.get('pharmacies'):
        add_colored_heading(doc, "Pharmacies", level=2, color_hex="5C7B9C")
        for ph in pratiques.get('pharmacies', []):
            add_info_box(doc, ph.get('nom', ''), ph.get('adresse', '') + (' — ' + ph.get('telephone', '') if ph.get('telephone') else ''), "💊")

    if pratiques.get('medecin'):
        add_colored_heading(doc, "Médecin / Hôpital", level=2, color_hex="5C7B9C")
        med = pratiques.get('medecin', {})
        add_info_box(doc, med.get('nom', ''), med.get('adresse', '') + (' — ' + med.get('telephone', '') if med.get('telephone') else ''), "🏥")

    if pratiques.get('transports'):
        add_colored_heading(doc, "Transports", level=2, color_hex="5C7B9C")
        for t in pratiques.get('transports', []):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.add_run("🚌 " + str(t)).font.size = Pt(10)
    add_divider(doc)

    # === URGENCES ===
    doc.add_paragraph()
    add_colored_heading(doc, "🚨  Numéros d'urgence", level=1)
    urg = guide_data.get('urgences', {})
    urgences_list = [
        ("SAMU", urg.get('samu', '15'), "🏥"),
        ("Pompiers", urg.get('pompiers', '18'), "🚒"),
        ("Police", urg.get('police', '17'), "🚔"),
        ("Urgences européen", urg.get('urgences_europeen', '112'), "🆘"),
    ]
    for label, num, emoji in urgences_list:
        add_info_box(doc, label, num, emoji)

    if urg.get('hopital_proche', {}).get('nom'):
        hop = urg.get('hopital_proche', {})
        doc.add_paragraph()
        add_colored_heading(doc, "Hôpital le plus proche", level=2, color_hex="5C7B9C")
        add_info_box(doc, hop.get('nom', ''), hop.get('adresse', ''), "🏨")
        if hop.get('telephone'):
            add_info_box(doc, "Téléphone", hop.get('telephone', ''), "📞")
    add_divider(doc)

    # === MOT DE FIN ===
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("✦  " + guide_data.get('mot_fin', 'Nous vous souhaitons un merveilleux séjour !') + "  ✦")
    fr.font.size = Pt(13)
    fr.font.italic = True
    fr.font.color.rgb = RGBColor(0xD4, 0xA9, 0x6A)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run("Guide généré avec Léon · leon-audit.fr")
    footer_r.font.size = Pt(9)
    footer_r.font.color.rgb = RGBColor(0xCC, 0xBB, 0xAA)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/")
def index():
    import os
    return send_file(os.path.join(os.path.dirname(__file__), 'templates', 'index.html'))


@app.route("/analyze", methods=["POST"])
def analyze():
    try:
        images = []
        for key in request.files:
            f = request.files[key]
            if f and f.filename:
                mt = f.content_type or "image/jpeg"
                img_data = encode_image(f)
                images.append((img_data, mt))
        if not images:
            return jsonify({"success": False, "error": "Aucune image reçue"}), 400
        vision = call_1_vision(images)
        scoring = call_2_scoring(vision)
        reco = call_3_recommendations(vision, scoring)
        return jsonify({"success": True, "vision": vision, "scoring": scoring, "recommendations": reco})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/save-email", methods=["POST"])
def save_email():
    try:
        data = request.get_json()
        email = data.get('email', '').strip()
        if not email or '@' not in email:
            return jsonify({"success": False, "error": "Email invalide"}), 400

        api_key = os.environ.get("BREVO_API_KEY")
        if not api_key:
            return jsonify({"success": False, "error": "API key manquante"}), 500

        import urllib.request
        payload = {
            "email": email,
            "listIds": [2],
            "updateEnabled": True,
            "attributes": {"SOURCE": "Leon Beta"}
        }
        req = urllib.request.Request(
            "https://api.brevo.com/v3/contacts",
            data=json.dumps(payload).encode('utf-8'),
            headers={
                "Content-Type": "application/json",
                "api-key": api_key
            },
            method="POST"
        )
        try:
            with urllib.request.urlopen(req) as resp:
                pass
        except urllib.error.HTTPError as e:
            body = e.read().decode()
            if e.code == 400 and "duplicate" in body.lower():
                pass  # Email déjà existant — on laisse passer
            else:
                return jsonify({"success": False, "error": f"Brevo error {e.code}"}), 500

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/generate-guide", methods=["POST"])
def generate_guide():
    try:
        vision_data = json.loads(request.form.get('vision_data', '{}'))
        form_data = {
            'wifi_nom': request.form.get('wifi_nom', ''),
            'wifi_mdp': request.form.get('wifi_mdp', ''),
            'parking': request.form.get('parking', ''),
            'code_acces': request.form.get('code_acces', ''),
            'infos_specifiques': request.form.get('infos_specifiques', '')
        }
        guide_data = call_guide(vision_data, form_data)
        buf = generate_guide_word(guide_data)
        titre = guide_data.get('titre_guide', 'Guide d\'accueil').replace('/', '-').replace('\\', '-')
        return send_file(buf, as_attachment=True, download_name=titre + '.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
